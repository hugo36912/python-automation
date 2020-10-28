from openpyxl import *
from docx import Document

def w3e():
        fil=''

        fil=input('filename :')
        sour=[]
        dest=[]
        app=[]
        pro=[]
        i=1
        try:
         dj = Document('Z:/Change_Request/2020/FWCR_2020/2020-eHR_FWCR_Production_Outstanding/'+fil+'.docx')
        except:
         dj = Document('Z:/Change_Request/2020/FWCR_2020/2020-PUB_FWCR_Production_Outstanding/'+fil+'.docx')
        kj=dj.tables[0]
        j=1
        k=1
        l=1
        descri=kj.cell(0,0).text.split(":")[1]
        descri=descri.replace('(','')
        descri=descri.replace(')','')
        descri=descri.replace(' ','')
        sourcol=0
        destcol=0
        appcol=0
        procol=0
        aj=1
        while  dj.tables[0].cell(i,0).text != '1.' and dj.tables[0].cell(i,0).text != '1':
            i+=1
        srow=i-1
        if "Source IP" not in dj.tables[0].cell(srow, 1).text:
            srow=i-2

        while aj < len(kj.columns):
            if  'Source IP' in dj.tables[0].cell(srow, aj).text:
                sourcol=aj
                break
            aj+=1
        aj=1
        while aj < len(kj.columns):
            if  'Destination IP' in dj.tables[0].cell(srow, aj).text:
                destcol=aj
                break
            aj+=1
        aj=1
        while aj < len(kj.columns):
            if "Protocol" in dj.tables[0].cell(srow, aj).text:
                procol=aj
                break
            aj+=1
        aj=1
        while aj < len(kj.columns):
            if  "Port Number" in dj.tables[0].cell(srow, aj).text :
                appcol=aj
                break
            aj+=1


        while i<len(kj.rows):
            ta = dj.tables[0].cell(i, 0).text
            if "Additional Information" in ta:
                break
            ta=ta.replace(".",'')
            ta = ta.replace("a", '')
            ta = ta.replace("b", '')
            ta = ta.replace("d", '')
            ta = ta.replace("e", '')
            ta = ta.replace("f", '')
            ta = ta.replace("g", '')
            ta = ta.replace("h", '')
            ta = ta.replace("i", '')
            if ta.isdigit():
                 sour.append(dj.tables[0].cell(i, sourcol).text)
                 dest.append(dj.tables[0].cell(i, destcol).text)
                 app.append(dj.tables[0].cell(i, appcol).text)
                 pro.append(dj.tables[0].cell(i, procol).text)
            i+=1
        findest=[]
        finsour=[]
        finapp=[]
        destre=[i.split('\n',66)for i in dest]
        for ib in destre:
            tmpdest = []
            for ic in ib:
                ic = ic.split('/', 1)[0]
                if '.' in ic :
                 tmpdest.append(ic)

            findest.append(tmpdest)

        sourr=[i.split('\n',66)for i in sour]

        for ib in sourr:
            tmpdest = []
            for ic in ib:
                 ic = ic.split('/', 1)[0]
                 if '.' in ic :
                  tmpdest.append(ic)

            finsour.append(tmpdest)

        appr = [i.split('\n', 66) for i in app]
        for ib in appr:
            tmpdest = []
            for ic in ib:
                if ic=='80':
                   ic="junos-http"
                if ic=='443':
                 ic="junos-https"
                if ic=='22':
                   ic="junos-ssh"
                if ic=='25':
                   ic="junos-smtp"
                if ic=='123':
                   ic="junos-ntp"
                if ic=='53':
                   ic="junos-dns-udp"
                if ic=='587':
                   ic="junos-smtps"
                tmpdest.append(ic)
            finapp.append(tmpdest)
        ih=0
        while ih < len(finsour):
          finsour[ih]=','.join(finsour[ih])
          finsour[ih]=finsour[ih].replace(",","\n")
          finsour[ih] = finsour[ih].replace("[", "")
          finsour[ih] = finsour[ih].replace("]", "")
          ih+=1

        ih=0
        while ih < len(findest):
          findest[ih]=','.join(findest[ih])
          findest[ih]=findest[ih].replace(",","\n")
          findest[ih] = findest[ih].replace("[", "")
          findest[ih] = findest[ih].replace("]", "")
          ih+=1
        ih = 0
        while ih < len(finapp):
            finapp[ih] = ','.join(finapp[ih])
            finapp[ih] = finapp[ih].replace(",", "\n")
            finapp[ih] = finapp[ih].replace("[", "")
            finapp[ih] = finapp[ih].replace("]", "")
            ih += 1
        ih = 0


        wb=load_workbook("Z:/Network_Security/Firewall_Policy_Program/PythonFirewall/PRDv3/connectDatabase/Database/eHR_PRD_FW_gpinput.xlsx")
        ws=wb["input_lookup"]
        for row in ws['A2':'J60']:
            for cell in row:
                cell.value=None
        r=2
        i=1
        while i< len(finsour)+1:
            ws.cell(r,1).value=i
            ws.cell(r, 2).value = finsour[i-1]
            ws.cell(r, 3).value = findest[i-1]
            ws.cell(r, 4).value = finapp[i-1]
            ws.cell(r, 5).value = pro[i-1]
            ws.cell(r, 6).value = descri
            r+=1
            i+=1

        wb.save("Z:/Network_Security/Firewall_Policy_Program/PythonFirewall/PRDv3/connectDatabase/Database/eHR_PRD_FW_gpinput.xlsx")
        print("success")
while True:
 w3e()